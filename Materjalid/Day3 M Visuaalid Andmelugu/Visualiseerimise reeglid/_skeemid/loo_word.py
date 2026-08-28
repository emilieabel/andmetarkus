# -*- coding: utf-8 -*-
"""Kujundatud Word-dokumentatsioon — Liikumisaktiivsuse programmid."""
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor, Emu

ROOT = Path(__file__).resolve().parent
DOC_DIR = ROOT.parent
SKEEM = ROOT
OUT = DOC_DIR / "Liikumisaktiivsuse_programmid_too_dokumentatsioon.docx"

TEAL = RGBColor(0x0A, 0x3D, 0x3A)
TEAL2 = RGBColor(0x1A, 0x3D, 0x40)
ACCENT = RGBColor(0x1B, 0x5F, 0x5A)
INK = RGBColor(0x2A, 0x2A, 0x2A)
MUTED = RGBColor(0x4A, 0x5C, 0x5A)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
FILL_TEAL = "0A3D3A"
FILL_HEAD = "1A3D40"
FILL_ROW = "E8F6F4"
FILL_ALT = "F5F5F5"
FILL_ACCENT = "2EC4B6"
FILL_CALL = "E8F6F4"


def set_run_font(run, name="Calibri", size=11, bold=False, italic=False, color=INK):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color


def shade(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def set_cell_margins(cell, top=40, bottom=40, left=80, right=80):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for m, v in (("top", top), ("left", left), ("bottom", bottom), ("right", right)):
        node = OxmlElement(f"w:{m}")
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")
        tcMar.append(node)
    tcPr.append(tcMar)


def set_table_borders(table, color="D0D8D7", sz="4"):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), sz)
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        borders.append(el)
    tblPr.append(borders)


def no_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "nil")
        borders.append(el)
    tblPr.append(borders)


def prevent_row_split(row):
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    cant = OxmlElement("w:cantSplit")
    trPr.append(cant)


def set_keep_with_next(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    k = OxmlElement("w:keepNext")
    pPr.append(k)


def cell_text(cell, text, *, size=10, bold=False, color=INK, align="left", fill=None):
    if fill:
        shade(cell, fill)
    set_cell_margins(cell)
    p = cell.paragraphs[0]
    p.text = ""
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "right":
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    return p


def add_table(doc, headers, rows, col_widths=None, header_fill=FILL_HEAD):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table, "C5D4D2", "4")
    usable = 16.6
    if col_widths is None:
        col_widths = [usable / len(headers)] * len(headers)
    for i, w in enumerate(col_widths):
        for cell in table.columns[i].cells:
            cell.width = Cm(w)
    for i, h in enumerate(headers):
        cell_text(table.rows[0].cells[i], h, size=10, bold=True, color=WHITE, fill=header_fill)
    prevent_row_split(table.rows[0])
    for r_i, row in enumerate(rows):
        fill = FILL_ROW if r_i % 2 == 0 else FILL_ALT
        prevent_row_split(table.rows[r_i + 1])
        for c_i, val in enumerate(row):
            cell_text(table.rows[r_i + 1].cells[c_i], str(val), size=10, fill=fill)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return table


def p(doc, text, *, size=11, bold=False, italic=False, color=INK, space_after=8, space_before=0, align="left"):
    para = doc.add_paragraph()
    if align == "center":
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "justify":
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    else:
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    para.paragraph_format.space_after = Pt(space_after)
    para.paragraph_format.space_before = Pt(space_before)
    para.paragraph_format.line_spacing = 1.15
    run = para.add_run(text)
    set_run_font(run, size=size, bold=bold, italic=italic, color=color)
    return para


def rich(doc, parts, *, space_after=8, justify=True):
    """parts: list of (text, bold, italic) or str."""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY if justify else WD_ALIGN_PARAGRAPH.LEFT
    para.paragraph_format.space_after = Pt(space_after)
    para.paragraph_format.line_spacing = 1.15
    for part in parts:
        if isinstance(part, str):
            run = para.add_run(part)
            set_run_font(run)
        else:
            text, bold, italic = part[0], part[1], part[2] if len(part) > 2 else False
            run = para.add_run(text)
            set_run_font(run, bold=bold, italic=italic)
    return para


def h1(doc, text):
    para = doc.add_paragraph(text, style="Heading 1")
    for run in para.runs:
        set_run_font(run, size=18, bold=True, color=TEAL)
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(10)
    set_keep_with_next(para)
    return para


def h2(doc, text):
    para = doc.add_paragraph(text, style="Heading 2")
    for run in para.runs:
        set_run_font(run, size=14, bold=True, color=ACCENT)
    para.paragraph_format.space_before = Pt(14)
    para.paragraph_format.space_after = Pt(6)
    set_keep_with_next(para)
    return para


def bullets(doc, items):
    for item in items:
        para = doc.add_paragraph(style="List Bullet")
        para.paragraph_format.space_after = Pt(3)
        para.paragraph_format.left_indent = Cm(0.75)
        run = para.add_run(item)
        set_run_font(run, size=11)


def numbered(doc, items):
    for item in items:
        para = doc.add_paragraph(style="List Number")
        para.paragraph_format.space_after = Pt(3)
        para.paragraph_format.left_indent = Cm(0.75)
        run = para.add_run(item)
        set_run_font(run, size=11)


def add_figure(doc, filename, caption, width=16.4):
    path = SKEEM / filename
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(8)
    para.paragraph_format.space_after = Pt(2)
    run = para.add_run()
    run.add_picture(str(path), width=Cm(width))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(12)
    cap.paragraph_format.space_before = Pt(2)
    r = cap.add_run(caption)
    set_run_font(r, size=9, italic=True, color=MUTED)


def callout(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table, "2EC4B6", "12")
    cell = table.cell(0, 0)
    shade(cell, FILL_CALL)
    set_cell_margins(cell, top=80, bottom=80, left=120, right=120)
    cell.width = Cm(16.6)
    p1 = cell.paragraphs[0]
    p1.paragraph_format.space_after = Pt(4)
    r1 = p1.add_run(title)
    set_run_font(r1, size=11, bold=True, color=TEAL)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(body)
    set_run_font(r2, size=10.5, color=TEAL2)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(8)


def page_break(doc):
    doc.add_page_break()


def add_header_footer(doc):
    section = doc.sections[0]
    section.different_first_page_header_footer = True
    # empty first page header/footer
    section.first_page_header.paragraphs[0].text = ""
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = fp.add_run("Töö dokumentatsioon  ·  Liikumisaktiivsuse programmid  ·  ")
    set_run_font(run, size=8, color=MUTED)
    # PAGE field
    fld1 = OxmlElement("w:fldChar")
    fld1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld2 = OxmlElement("w:fldChar")
    fld2.set(qn("w:fldCharType"), "end")
    r2 = fp.add_run()
    set_run_font(r2, size=8, color=MUTED)
    r2._r.append(fld1)
    r2._r.append(instr)
    r2._r.append(fld2)

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = hp.add_run("Liikumisaktiivsuse programmid")
    set_run_font(r, size=9, bold=True, color=TEAL)
    r3 = hp.add_run("   ·   Power BI töö dokumentatsioon")
    set_run_font(r3, size=9, color=MUTED)
    # bottom border on header
    pPr = hp._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "2EC4B6")
    pBdr.append(bottom)
    pPr.append(pBdr)


def style_base(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    pf = normal.paragraph_format
    pf.space_after = Pt(8)
    pf.line_spacing = 1.15

    for name, size, color in (("Heading 1", 18, TEAL), ("Heading 2", 14, ACCENT), ("Heading 3", 12, TEAL2)):
        st = styles[name]
        st.font.name = "Calibri"
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = color
        st.font.color.theme_color = None


def cover(doc):
    # top banner
    t = doc.add_table(1, 1)
    no_table_borders(t)
    c = t.cell(0, 0)
    shade(c, FILL_TEAL)
    set_cell_margins(c, top=280, bottom=280, left=200, right=200)
    p1 = c.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p1.add_run("ANDMETARKUS  ·  GRUPI TÖÖ")
    set_run_font(r, size=11, bold=True, color=RGBColor(0x2E, 0xC4, 0xB6))
    p2 = c.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p2.paragraph_format.space_before = Pt(10)
    r = p2.add_run("Liikumisaktiivsuse programmid")
    set_run_font(r, size=28, bold=True, color=WHITE)
    p3 = c.add_paragraph()
    r = p3.add_run("Power BI aruande töö dokumentatsioon")
    set_run_font(r, size=14, color=RGBColor(0xD4, 0xF4, 0xF0))

    p(doc, "", space_after=6)
    p(doc, "Avaliku liikumisaktiivsuse tegevuskava seire visualiseerimine", size=13, color=TEAL, space_after=16)

    meta = [
        ("Aruanne", "Liikumisaktiivsuse programmid.pbip"),
        ("Andmeallikas", "app.liigume.ee/api  (activities, organizations, policies)"),
        ("Andmete seis", "28.08.2026  ·  100 tegevust"),
        ("Tööriist", "Microsoft Power BI  ·  PBIP / TMDL  ·  et-EE"),
        ("Organisatsioon", "SA Liikumisharrastuse kompetentsikeskus (LHKK)"),
        ("Grupiliikmed", "………………………………………………………………"),
    ]
    add_table(doc, ["Väli", "Väärtus"], meta, col_widths=[4.2, 12.4])

    callout(
        doc,
        "Peamine sõnum",
        "Kava on 2026. aasta suveks elus koostööportfell, mis ehitab peamiselt süsteeme — juhtimist, andmeid, teadust ja taristut. Inimeste hoiakute suund on kõige õhem ja edenemise protsent puudub 72 real.",
    )

    p(doc, "Sisukord", size=14, bold=True, color=TEAL, space_before=8, space_after=6)
    toc = [
        ("1.", "Ettevõtte ja uurimisprobleemi tutvustus"),
        ("2.", "Grupitöö plaan"),
        ("3.", "Ärisõnastik"),
        ("4.", "Andmekaitse kirjeldus"),
        ("5.", "Andmemudel, andmesõnastik"),
        ("6.", "Andmeallikad ja andmevoog"),
        ("7.", "Näidisandmestiku loomine, hankimine, import"),
        ("8.", "Andmete kvaliteedi kontroll"),
        ("9.", "Andmete töötlemine"),
        ("10.", "Andmete analüüs"),
        ("11.", "Kirjeldav raport / analüüs"),
        ("12.", "Andmelugu, järeldused"),
    ]
    add_table(doc, ["Pt", "Peatükk"], toc, col_widths=[1.6, 15.0])
    p(
        doc,
        "Joonised: tähemudel, andmevoog, grupitöö etapid, WHO suunad, aruande lehed, andmekvaliteet, staatus.",
        size=9,
        italic=True,
        color=MUTED,
        space_after=0,
    )


def build():
    doc = Document()
    style_base(doc)
    sec = doc.sections[0]
    sec.page_width = Cm(21.0)
    sec.page_height = Cm(29.7)
    sec.left_margin = Cm(2.1)
    sec.right_margin = Cm(2.1)
    sec.top_margin = Cm(2.2)
    sec.bottom_margin = Cm(2.0)
    sec.header_distance = Cm(1.0)
    sec.footer_distance = Cm(1.0)
    add_header_footer(doc)

    cover(doc)

    # ----- 1 -----
    page_break(doc)
    h1(doc, "1. Ettevõtte ja uurimisprobleemi tutvustus")
    h2(doc, "1.1 Organisatsioon")
    rich(
        doc,
        [
            "Andmestiku omanik ja tegevuskava seire platvormi haldaja on ",
            ("SA Liikumisharrastuse kompetentsikeskus", True, False),
            " (LHKK, liigume.ee). Keskuse asutas 2022. aasta sügisel Eesti Olümpiakomitee. Missioon on olla liikumispöörde eestvedaja, visioon aastaks 2035 — ",
            ("Euroopa liikuvaim rahvas", True, False),
            ".",
        ],
    )
    rich(
        doc,
        [
            "Tegevuskava ei ole ühe asutuse siseprojekt. 2022. aasta sügisel algatasid kultuuri-, sotsiaal- ning haridus- ja teadusministrid valdkondadeülese koostöö; koostamist vedas ",
            ("Kultuuriministeerium", True, False),
            ". 2023. aastal valmis ühine liikumisaktiivsuse tegevuskava, mis tugineb arengustrateegiale „Eesti 2035“ ja spordipoliitika alusdokumendile ",
            ("Sport 2030", True, False),
            " (eesmärk: aastaks 2030 liigub ja spordib vähemalt kaks kolmandikku elanikest). Aastatel 2024–2025 liitusid kava kaitse- ja transpordivaldkond, sh Eesti rattastrateegia 2040.",
        ],
    )
    rich(
        doc,
        [
            "Struktuurilt järgib kava Maailma Terviseorganisatsiooni ülemaailmset kehalise aktiivsuse tegevuskava ",
            ("WHO GAPPA 2018–2030", True, False),
            " nelja strateegilist suunda ja 20 poliitikameedet.",
        ],
    )

    h2(doc, "1.2 Miks see teema on oluline")
    p(
        doc,
        "Liikumisaktiivsus Eestis on alla tervisesoovituste. WHO soovitab täiskasvanutel liikuda vähemalt 2,5–5 tundi nädalas mõõduka intensiivsusega. Enamik elanikkonnast seda ei täida; madal kehaline aktiivsus suurendab krooniliste haiguste riski ja lühendab tervena elatud aastaid.",
        align="justify",
    )
    add_table(
        doc,
        ["Näitaja", "Väärtus", "Allikas"],
        [
            ["Täiskasvanud, kes tegelevad liikumise/spordiga > 1× nädalas", "44%", "liigume.ee/tegevuskava"],
            ["15-aastased, kehaliselt aktiivsed ≥ 5 päeval nädalas", "37%", "samas"],
            ["Vanemaealised (55+), aktiivsed > 2 päeval nädalas", "19%", "samas"],
            ["4. klassi lapsed, kes on ülekaalulised või rasvunud", "34%", "samas"],
            ["Ajateenijad, kes sooritavad KV testi teenistuse alguses positiivselt", "31%", "samas"],
            ["Tervena elada jäänud aastad (2024)", "mehed 56,8; naised 60,6", "Statistikaamet"],
        ],
        col_widths=[8.4, 3.6, 4.6],
    )

    h2(doc, "1.3 Uurimisprobleem")
    rich(
        doc,
        [
            "Tegevuskava on rulluv dokument (100 avalikku tegevust, periood kuni 2030) ja seda seiratakse veebirakenduses app.liigume.ee. Avalik vaade näitab tegevusi kaardina, kuid ",
            ("ei anna analüütikule koondpilti", True, False),
            ": kui suur osa kavast on tehtud, milline WHO suund on üle- või alakaetud, kes juhib elluviimist, kui hästi on täidetud edenemine ja indikaatorid ning kus on andmeaugud.",
        ],
    )
    p(doc, "Uurimisküsimused", bold=True, size=11, color=TEAL, space_after=4)
    numbered(
        doc,
        [
            "Kui kaugel on 100 tegevuse portfell staatuse, tüübi ja ajaperioodi järgi?",
            "Kuidas jaguneb vastutus asutuste ja WHO GAPPA suundade vahel?",
            "Kas kõik 20 poliitikameedet on tegevustega kaetud ja millised sihtrühmad on fookuses?",
            "Kui usaldusväärne on seireandmestik (kuupäevad, edenemise %, indikaatorid, partnerid)?",
        ],
    )
    rich(
        doc,
        [
            ("Töö eesmärk. ", True, False),
            "Koostada Power BI aruanne, mis muudab avaliku API andmed juhtimislaadseks ülevaateks: mahud, võrgustik, poliitikakate ja andmekvaliteet ühes mudelis.",
        ],
    )
    rich(
        doc,
        [
            ("Sihttarbijad. ", True, False),
            "Tegevuskava koordinatsioon (LHKK, Kultuuriministeerium), ministeeriumide ja ametite kontaktisikud, kursuse hindajad. Aruanne on õppetöö analüütiline prototüüp, mitte LHKK ametlik seirearuanne.",
        ],
    )

    # ----- 2 -----
    page_break(doc)
    h1(doc, "2. Grupitöö plaan")
    p(
        doc,
        "Töö jagunes neljaks etapiks, mis vastavad andmeprojekti elutsüklile. Iga etapp lõppes kontrollitava tulemiga (Power Query päring, TMDL-mudel, aruande leht või see dokumentatsioon).",
        align="justify",
    )
    add_figure(doc, "skeem_grupitoo.png", "Joonis 3. Grupitöö neli etappi")
    h2(doc, "2.1 Tööjaotus")
    add_table(
        doc,
        ["Etapp", "Sisu", "Tulem"],
        [
            ["A. Probleem ja mõisted", "Taust, uurimisküsimused, ärisõnastik, andmekaitse", "ptk 1–4"],
            ["B. Andmed ja mudel", "API, Power Query, tähemudel, mõõdikud", "ptk 5–9, Semantic Model"],
            ["C. Visuaalid", "Viis aruandelehte, teema, filtrid, Info-leht", "Report"],
            ["D. Analüüs ja lugu", "Kirjeldav statistika, andmelugu, järeldused", "ptk 10–12"],
        ],
        col_widths=[4.2, 8.2, 4.2],
    )
    h2(doc, "2.2 Ajagraafik")
    add_table(
        doc,
        ["Jr", "Tegevus", "Sõltuvus"],
        [
            ["1", "API otspunktide kaardistamine", "—"],
            ["2", "Toorandmete laadimine ja skeemi kirjeldamine", "1"],
            ["3", "Kvaliteedikontroll (kuupäevad, edenemine, mitu-mitmele)", "2"],
            ["4", "Tähemudeli koostamine (fakt + 3 dimi + mõõdikud)", "3"],
            ["5", "DAX-mõõdikud viies kaustas", "4"],
            ["6", "Aruandelehed Ülevaade → … → Info", "5"],
            ["7", "INFO.VIEW kataloog Info-lehel", "5–6"],
            ["8", "Kirjeldav analüüs ja andmelugu", "6"],
        ],
        col_widths=[1.4, 12.4, 2.8],
    )
    h2(doc, "2.3 Kokkulepped")
    bullets(
        doc,
        [
            "Power BI Desktop (DevMode, PBIP), kultuur et-EE.",
            "Import-režiim, värskendus käsitsi (Home → Refresh), autentimine Anonymous.",
            "Tarbijanimed eesti keeles, tehnilised võtmed inglise keeles (activity_id, lead_org_id).",
            "Mitu-mitmele seoseid ei modelleerita: partnerid, meetmed, sihtrühmad ja indikaatorid jäävad tegevuse reale.",
            "Implicit measures on keelatud — kõik koondarvud on explitsiitsed DAX-mõõdikud.",
        ],
    )
    h2(doc, "2.4 Riskid")
    add_table(
        doc,
        ["Risk", "Mõju", "Leevendus"],
        [
            ["API muutub või on kättesaamatu", "Aruanne ei värskene", "Dokumenteeritud otspunktid; stabiilne JSON"],
            ["Edenemise % puudub enamikul ridadel", "KPI on kallutatud", "Eraldi mõõdik; keskmine ainult täidetud ridadel"],
            ["Partnerid/meetmed ühel tekstiväljal", "Unikaalset partnerit ei saa filtrerida", "Arvuveerud ja selged visuaalipealkirjad"],
            ["Grupiliikmete nimed pole fikseeritud", "Rollijaotus hindamisel", "Lisada nimed esilehele enne esitamist"],
        ],
        col_widths=[5.4, 4.4, 6.8],
    )

    # ----- 3 -----
    page_break(doc)
    h1(doc, "3. Ärisõnastik")
    p(doc, "Mõisted, mida aruandes ja mudelis kasutatakse. Tehniline andmesõnastik on peatükis 5.", align="justify")
    add_table(
        doc,
        ["Mõiste", "Tähendus selles töös"],
        [
            ["Liikumisaktiivsuse tegevuskava", "Valdkondadeülene rulluv riiklik kava kehalise aktiivsuse tõstmiseks (alates 2023, seire kuni 2030)."],
            ["Tegevus (activity)", "Üks kirje kavas: projekt, programm, uuring, kampaania või õigusruumi muudatus. Mudelis üks rida fact_activity."],
            ["Tegevuse tüüp", "Eesti projektid ja programmid  või  Teadus ja innovatsioon."],
            ["Staatus", "Pole alustatud, Töös, Tehtud. Järjestus 1–3."],
            ["Edenemine %", "Valmidus 0–100. Paljudel ridadel puudub."],
            ["Indikaator", "Alameesmärk nime ja staatusega (Ootel, Pooleli, Lõpetatud, Mõõdik)."],
            ["Sihttulemus", "Lühike kirjeldus, mida tegevuse lõpuks oodatakse."],
            ["Periood", "Ajavahemik tekstina (nt 2025–2027) või väärtus Pidev."],
            ["Juhtorganisatsioon", "Asutus, kes vastutab elluviimise eest (lead_org_id)."],
            ["Partner", "Kaasatud asutus. API-s loend; mudelis komadega tekst."],
            ["Organisatsiooni tüüp", "Ministeerium, Riigiamet, Riigiasutus, MTÜ, Ülikool, KOV, Muu."],
            ["WHO suund (GAPPA)", "Neli strateegilist suunda: ühiskond, keskkond, inimene, struktuurid."],
            ["Poliitikameede", "GAPPA 20 meedet (koodid 1.1–4.5). Tegevusel võib olla mitu."],
            ["Sihtrühm", "Lapsed ja noored; Elukaareülene; Keskkond; Teadus ja innovatsioon."],
            ["Partnerlus", "Üks seos tegevuse ja partnerasutuse vahel; mõõdik liidab need kokku."],
            ["Pidev tegevus", "Algatus ilma kindla lõputa; tihti puuduvad kuupäevad."],
            ["Tähtaeg ületatud", "Lõppkuupäev on möödas, kuid staatus ei ole Tehtud."],
            ["LHKK", "SA Liikumisharrastuse kompetentsikeskus."],
            ["Sport 2030", "Spordipoliitika alusdokument; siht kahele kolmandikule elanikest."],
        ],
        col_widths=[4.6, 12.0],
    )
    h2(doc, "WHO GAPPA neli suunda")
    add_table(
        doc,
        ["Kood", "Suund mudelis", "Fookus"],
        [
            ["1", "Loome aktiivset ühiskonda", "Normid, teadlikkus, kampaaniad, hoiakud"],
            ["2", "Loome aktiivset keskkonda", "Ruum, taristu, liiklusohutus, avalik ruum"],
            ["3", "Loome aktiivset inimest", "Programmid, haridus, tervis, võimalused"],
            ["4", "Loome aktiivseid struktuure", "Juhtimine, andmed, teadus, rahastus"],
        ],
        col_widths=[2.0, 6.4, 8.2],
    )
    add_figure(doc, "skeem_who_suunad.png", "Joonis 4. WHO GAPPA suunad tegevuste arvu järgi (n = 100)")

    # ----- 4 -----
    page_break(doc)
    h1(doc, "4. Andmekaitse kirjeldus")
    h2(doc, "4.1 Milliseid andmeid töödeldakse")
    rich(
        doc,
        [
            "Töödeldakse ",
            ("avalikke haldusandmeid", True, False),
            " tegevuskava kohta: asutuste nimed, tegevuste pealkirjad, staatused, kuupäevad, poliitikameetmed, sihtrühmade sildid. Andmestikus ",
            ("ei ole", True, False),
            " füüsiliste isikute nimesid, isikukoode, kontaktandmeid, terviseandmeid ega muid IKÜM art 9 eriliigilisi isikuandmeid. API on avalik ja autentimata (Anonymous). Sama sisu on nähtav veebis app.liigume.ee.",
        ],
    )
    h2(doc, "4.2 Kas tegemist on isikuandmetega?")
    rich(
        doc,
        [
            "Isikuandmete kaitse üldmäärus (IKÜM / GDPR) kohaldub isikuandmetele (art 2 ja 4). Asutuste nimed, lühendid ja tüübid ei ole isikuandmed. Tegevuste kirjeldused on poliitika- ja programmitasandi tekstid.",
        ],
    )
    callout(
        doc,
        "Järeldus",
        "Käesoleva aruande andmestik ei ole isikuandmete töötlemine IKÜM mõttes. Õiguslik raamistik on siiski kirjas, sest tegemist on avaliku sektori teabe ja õppetöö dokumendiga.",
    )
    h2(doc, "4.3 Õiguslikud alused")
    add_table(
        doc,
        ["Alus", "Rakendus siin"],
        [
            ["Avaliku teabe seadus (AvTS)", "Tegevuskava on avalik teave. Analüüs kasutab juba avalikustatud teavet."],
            ["IKÜM art 6 lg 1 p e", "Avaliku ülesande täitmine on teabevaldaja, mitte üliõpilasrühma alus."],
            ["IKÜM art 6 lg 1 p f", "Õigustatud huvi õppetööks, kui isikuandmeid siiski esineks; andmed on juba avalikud."],
            ["IKÜM art 89 / õppetöö", "Ainult kursuse analüüs ja aruanne, mitte profiilianalüüs ega turundus."],
            ["Andmete minimeerimine (art 5)", "Imporditakse analüüsiks vajalikud väljad; tehnilised võtmed on peidetud."],
        ],
        col_widths=[5.0, 11.6],
    )
    h2(doc, "4.4 Töötlemise põhimõtted")
    numbered(
        doc,
        [
            "Eesmärgipiirang — tegevuskava seire visualiseerimine õppetöös.",
            "Säilitamine — koopia elab kohalikus .pbip projektis. Power BI Service’isse avaldamist ei nõuta.",
            "Edastamine — kolmandatele isikutele isikuandmeid ei edastata. API päringud TLS-ühendusega.",
            "Turvalisus — autentimist ei kasutata, sest allikas on avalik. Salajasi võtmeid projektis ei ole.",
            "Eriliigilisi andmeid ei ole — research_ids on uuringuviited, mitte uuritavate kirjed.",
            "Vastutav töötleja avalikus ruumis on teabevaldaja (LHKK / ministeeriumid). Üliõpilased on andmete kasutajad.",
        ],
    )

    # ----- 5 -----
    page_break(doc)
    h1(doc, "5. Andmemudel, andmesõnastik")
    h2(doc, "5.1 Mudeli tüüp")
    rich(
        doc,
        [
            "Semantiline mudel on ",
            ("tähemudel (star schema)", True, False),
            ": keskel fakt fact_activity (üks rida = üks tegevus), ümber kolm dimensiooni. Auto Date/Time on välja lülitatud; kalender on eraldi tabel CALENDAR(2023-01-01, 2030-12-31).",
        ],
    )
    add_figure(doc, "skeem_tahemodel.png", "Joonis 1. Semantiline mudel — tähemudel")
    add_table(
        doc,
        ["Seos", "From", "To", "Aktiivne", "Kard."],
        [
            ["fact_activity_to_dim_direction", "who_direction_id", "direction_id", "jah", "m : 1"],
            ["fact_activity_to_dim_organization", "lead_org_id", "organization_id", "jah", "m : 1"],
            ["fact_activity_to_dim_date", "Alguskuupäev", "Kuupäev", "jah", "m : 1"],
            ["fact_activity_end_to_dim_date", "Lõppkuupäev", "Kuupäev", "ei", "m : 1"],
        ],
        col_widths=[5.6, 3.6, 3.2, 2.2, 2.0],
    )
    callout(
        doc,
        "Teadlik lihtsustamine",
        "Partnerid, poliitikameetmed, sihtrühmad ja indikaatorid ei ole sildtabelid. Need on tegevuse real denormaliseeritud tekst (komad / semikoolonid) plus loendurid. Unikaalset partnerit või meedet ei saa viilutada nagu dimensiooni — aruande pealkirjad ütlevad selle kasutajale otse.",
    )
    h2(doc, "5.2 Tabelid")
    add_table(
        doc,
        ["Tabel", "Roll", "Ridu", "Partitsioon"],
        [
            ["fact_activity", "Fakt: tegevused", "100", "Power Query, import"],
            ["dim_organization", "Juhtorganisatsioonid", "51", "Power Query, import"],
            ["dim_direction", "WHO suunad (is_direction = 1)", "4", "Power Query, import"],
            ["dim_date", "Kalender", "2023–2030", "DAX CALENDAR"],
            ["_Measures", "Kõik DAX-mõõdikud", "1 dummy-rida", "arvutatud"],
            ["dim_measure_info", "Mõõdikute kataloog", "INFO.VIEW.MEASURES", "arvutatud"],
            ["dim_mudeli_objekt", "Tabelite/veergude kirjeldused", "INFO.VIEW", "arvutatud"],
        ],
        col_widths=[4.0, 5.2, 3.8, 3.6],
    )
    h2(doc, "5.3 Andmesõnastik — fact_activity")
    add_table(
        doc,
        ["Veerg", "Tüüp", "Kirjeldus"],
        [
            ["activity_id", "täisarv, PK, peidetud", "Tegevuse identifikaator"],
            ["lead_org_id", "täisarv, peidetud", "Juhtorganisatsiooni FK"],
            ["who_direction_id", "täisarv, peidetud", "WHO suuna FK"],
            ["Tegevuse nr / Kood / Tegevus / Lühinimi", "täisarv / tekst", "Identiteet kavas"],
            ["Tegevuse tüüp", "tekst", "Projekt/programm või teadus"],
            ["Periood", "tekst", "Ajavahemik või Pidev"],
            ["Alguskuupäev / Lõppkuupäev", "kuupäev", "Aktiivne / mitteaktiivne seos dim_date"],
            ["Staatus / Staatuse järjekord", "tekst / täisarv", "Pole alustatud, Töös, Tehtud (1–3)"],
            ["Sihttulemus", "tekst", "Oodatav tulemus"],
            ["Juhtorganisatsioon / lühend / tüüp", "tekst", "Denormaliseeritud juhtasutus"],
            ["Partnerid / Poliitikameetmed / Sihtrühmad", "tekst", "Mitu väärtust, komadega"],
            ["Indikaatorid", "tekst", "nimi (staatus), semikooloniga"],
            ["Edenemine protsent", "täisarv", "0–100, sageli tühi"],
            ["Partnerite / Meetmete / Indikaatorite arv", "täisarv, peidetud", "Loendite pikkused"],
            ["Lõpetatud indikaatorite arv", "täisarv, peidetud", "status = Lõpetatud"],
            ["Uuringute / Sihtrühmade arv", "täisarv, peidetud", "research_ids / target_groups"],
            ["Kestus päevades", "täisarv", "Lõpp − algus, kui mõlemad olemas"],
        ],
        col_widths=[5.4, 4.2, 7.0],
    )
    p(doc, "Hierarhia: Tegevuse tüüp ja staatus.", italic=True, size=10, color=MUTED)
    h2(doc, "5.4 Dimensioonid")
    p(doc, "dim_organization: organization_id, Organisatsioon, Lühend, Organisatsiooni tüüp, Peaorganisatsioon (self-join). Hierarhia: Organisatsiooni tüüp → Organisatsioon.", align="justify")
    p(doc, "dim_direction: direction_id, Suuna kood (1–4), Suund, Kirjeldus.", align="justify")
    p(doc, "dim_date: Kuupäev, Aasta, Kuu, Kvartal, Aasta-kuu. Hierarhia: Aasta → Kvartal → Kuu.", align="justify")
    h2(doc, "5.5 Mõõdikud")
    add_table(
        doc,
        ["Kaust", "Mõõdik", "Loogika"],
        [
            ["1. Ülevaade", "Tegevuste arv", "COUNTROWS(fact_activity)"],
            ["1. Ülevaade", "Tehtud / Töös / Alustamata", "CALCULATE staatuse filtriga"],
            ["1. Ülevaade", "Tehtud %", "Tehtud / kõik"],
            ["2. Edenemine", "Keskmine edenemine %", "AVERAGE(Edenemine) / 100 — ainult täidetud read"],
            ["2. Edenemine", "Indikaatorite arv / Lõpetatud / %", "SUM ja jagatis"],
            ["2. Edenemine", "Keskmine kestus päevades", "AVERAGE(Kestus päevades)"],
            ["3. Võrgustik", "Juhtorganisatsioonide arv", "DISTINCTCOUNT(lead_org_id)"],
            ["3. Võrgustik", "Partnerluste arv", "SUM(Partnerite arv)"],
            ["4. Poliitika", "Poliitikameetmete arv", "SUM(Meetmete arv) — kirjeid, mitte unikaalseid"],
            ["4. Poliitika", "Sihtrühmade arv", "SUM(Sihtrühmade arv)"],
            ["5. Andmekvaliteet", "Ilma partnerita / kuupäevata / edenemiseta", "CALCULATE + ISBLANK"],
            ["5. Andmekvaliteet", "Tähtaeg ületatud", "lõpp < TÄNA ja staatus ≠ Tehtud"],
        ],
        col_widths=[3.8, 5.6, 7.2],
    )

    # ----- 6 -----
    page_break(doc)
    h1(doc, "6. Andmeallikad ja andmevoog")
    h2(doc, "6.1 Allikad")
    p(doc, "Kolm REST JSON otspunkti, Power Query shared expressions. Lisaks DAX-kalender, mitte välisest failist.", align="justify")
    add_table(
        doc,
        ["Expression", "URL", "Sisaldus"],
        [
            ["Activities Source", "https://app.liigume.ee/api/activities", "100 tegevust (pesad loendid)"],
            ["Organizations Source", "https://app.liigume.ee/api/organizations", "51 asutust"],
            ["Policies Source", "https://app.liigume.ee/api/policies", "24 rida: 4 suunda + 20 meedet"],
        ],
        col_widths=[4.2, 8.2, 4.2],
    )
    add_figure(doc, "skeem_andmevoog.png", "Joonis 2. Andmeallikad ja andmevoog")
    h2(doc, "6.2 Värskendamine")
    p(
        doc,
        "Import, mitte DirectQuery. Iga värskendus laadib kogu JSON uuesti. Power BI Desktopis: Home → Refresh. Organisatsioonid ja poliitikad puhverdatakse (Table.Buffer) enne tegevuste lookup’e.",
        align="justify",
    )
    h2(doc, "6.3 Toore tegevuse väljad")
    p(
        doc,
        "id, nr, code, name, name_short, activity_type, period, start_date, end_date, status, progress_pct, target_outcome, description, target_groups[], who_direction_id, lead_org_id, partner_ids[], policy_ids[], research_ids[], indicators[{id, name, status}], read_more_link, last_modified_time.",
        size=10,
        italic=True,
        color=MUTED,
    )

    # ----- 7 -----
    page_break(doc)
    h1(doc, "7. Näidisandmestiku loomine, hankimine, import")
    rich(
        doc,
        [
            "Andmestikku ei genereeritud käsitsi ega Excelis. See on ",
            ("tootmislähedane avalik API", True, False),
            ", mida LHKK kasutab tegevuskava rakenduses. Õppetöö jaoks on see näidisandmestik: piiratud maht (100 rida), stabiilne skeem, puuduvad isikuandmed.",
        ],
    )
    h2(doc, "7.1 Hankimise sammud")
    numbered(
        doc,
        [
            "Kontrollida otspunkte brauseris või GET päringuga.",
            "Power BI-s Get data → Web (või olemasolevad expressions) aadressidele /api/activities, /api/organizations, /api/policies.",
            "Autentimine: Anonymous.",
            "Json.Document → Table.FromRecords.",
        ],
    )
    h2(doc, "7.2 Import")
    bullets(
        doc,
        [
            "Failivorming PBIP (kaustad .Report ja .SemanticModel), mitte üks .pbix.",
            "Režiim Import.",
            "Päringute järjekord: fact_activity, dim_organization, dim_direction, seejärel kolm Source expression’it.",
        ],
    )
    callout(
        doc,
        "Miks mitte sünteetiline näidis?",
        "Sünteetilised andmed ei paljastaks tegelikke seireauke (72 rida ilma edenemiseta, 38 ilma alguskuupäevata, meetmete ebaühtlane kate). Uurimisprobleem on just päris kava seiret lugeda.",
    )

    # ----- 8 -----
    page_break(doc)
    h1(doc, "8. Andmete kvaliteedi kontroll")
    p(
        doc,
        "Kontroll tehti toor-JSON-il (28.08.2026) ja on mudelis korduvkasutatav nelja andmekvaliteedi mõõdikuga. Aruande leht Edenemine näitab neid KPI-dena.",
        align="justify",
    )
    add_figure(doc, "skeem_kvaliteet.png", "Joonis 6. Andmekvaliteedi neli signaali (n = 100)")
    h2(doc, "8.1 Täielikkus")
    add_table(
        doc,
        ["Kontroll", "Tulemus (n = 100)", "Märkus"],
        [
            ["Staatus / tüüp / suund / juht / sihtrühm / meede / indikaator", "100%", "Kohustuslikud väljad on terved"],
            ["Unikaalseid juhtorganisatsioone", "22", "Registris 51 asutust"],
            ["Alguskuupäev puudub", "38", "Kõik 38 on perioodiga Pidev"],
            ["Lõppkuupäev puudub", "41", "Seotud pidevate tegevustega"],
            ["Edenemine % puudub", "72", "Täidetud ainult 28 real"],
            ["Partnereid pole", "22", "partner_ids tühi loend"],
            ["Tähtaeg ületatud", "0", "Möödunud lõpuga ridu, mis pole Tehtud, ei ole"],
        ],
        col_widths=[7.4, 3.6, 5.6],
    )
    h2(doc, "8.2 Vastuolud")
    bullets(
        doc,
        [
            "Perioodi tekst ei ole ühtne (en-kriips vs sidekriips; üks väärtus on ainult „2026“).",
            "41 tegevust on Pidev; neist 38-l puudub alguskuupäev. Kalendrivaated jätavad need vahele.",
            "19 Tehtud rida on kõik edenemisega 100. 58-st Töös tegevusest enamikul edenemist pole.",
            "Partner, meede, sihtrühm ja indikaator ei ole 1NF — unikaalset meedet ei saa usaldusväärselt tekstifiltriga lugeda.",
            "Mõõdik „Poliitikameetmete arv“ on kirjeite summa (105), mitte 17 unikaalset meedet.",
        ],
    )
    h2(doc, "8.3 Viidete terviklikkus")
    p(
        doc,
        "Kõik lead_org_id ja who_direction_id väärtused leiti vastavatest allikatest. partner_ids ja policy_ids lookup kasutab Record.FieldOrDefault — tundmatu id ei katkesta värskendust.",
        align="justify",
    )
    callout(
        doc,
        "Kvaliteedi järeldus",
        "Andmestik on struktuurselt korras (võtmed, staatused, suunad) ja sisuliselt auklik seireväljadel (kuupäev, edenemine). Aruanne ei peida auke: need on eraldi KPI-d.",
    )

    # ----- 9 -----
    page_break(doc)
    h1(doc, "9. Andmete töötlemine")
    p(doc, "Töötlus toimub Power Query M-keeles importimisel.", align="justify")
    h2(doc, "9.1 Ühendamised")
    numbered(
        doc,
        [
            "Tegevused ⟕ organisatsioonid lead_org_id = id → Juhtorganisatsioon, lühend, tüüp.",
            "partner_ids → nimed läbi organisatsioonide kaardi → Partnerid.",
            "policy_ids → policy_code + policy_measure → Poliitikameetmed.",
            "target_groups loend → unikaalsed, komadega Sihtrühmad.",
            "indicators objektid → nimi (staatus), semikooloniga Indikaatorid.",
            "dim_organization: self-join parent_organization_id → Peaorganisatsioon.",
            "dim_direction: policies read, kus is_direction = 1.",
        ],
    )
    h2(doc, "9.2 Tuletatud veerud")
    add_table(
        doc,
        ["Veerg", "Reegel"],
        [
            ["Partnerite / Meetmete / Indikaatorite / Uuringute / Sihtrühmade arv", "Vastava loendi pikkus (null = 0)"],
            ["Lõpetatud indikaatorite arv", "Indikaatorid, mille status = Lõpetatud"],
            ["Kestus päevades", "Duration.Days(end − start), kui mõlemad olemas"],
            ["Staatuse järjekord", "Pole alustatud = 1, Töös = 2, Tehtud = 3, muu = 9"],
        ],
        col_widths=[7.2, 9.4],
    )
    h2(doc, "9.3 Mida ei tehtud")
    bullets(
        doc,
        [
            "Ridu ei eemaldatud — 100 tegevust jäävad kõik alles.",
            "Puuduvaid kuupäevi ega edenemist ei täidetud vaikeväärtusega.",
            "Mitu-mitmele ei normaliseeritud sildtabeliteks.",
            "Perioodi kirjapilku allikas ei parandatud (jääb kvaliteedileiuks).",
        ],
    )
    h2(doc, "9.4 DAX pärast importi")
    p(
        doc,
        "Kalendritabel, mõõdikud filtrikontekstiga (CALCULATE, DIVIDE) ning INFO.VIEW.MEASURES / TABLES / COLUMNS — mudeli dokumentatsioon aruande sees.",
        align="justify",
    )

    # ----- 10 -----
    page_break(doc)
    h1(doc, "10. Andmete analüüs")
    rich(
        doc,
        [
            "Analüüs on ",
            ("kirjeldav", True, False),
            " (sagedused, osakaalud, katvus). Põhjuse-tagajärje mudelit ega regressiooni ei ehitata: n = 100 on kogu avalik kava, mitte valim rahvastikust. Arvud on API seisuga 28.08.2026.",
        ],
    )
    h2(doc, "10.1 Mida mõõdetakse")
    add_table(
        doc,
        ["Küsimus", "Meetod", "Väli / mõõdik"],
        [
            ["Kui valmis on portfell?", "Loendused, osakaal", "Staatus, Tehtud %"],
            ["Kas teadus vs programmid erinevad?", "Risttabel", "Tüüp × staatus"],
            ["Milline GAPPA suund kannab koormust?", "Loendus", "dim_direction"],
            ["Kes juhib?", "Loendus", "Juhtorganisatsioon"],
            ["Kas kõik 20 meedet on kasutusel?", "Unikaalsed policy_ids", "Võrdlus policies-loendiga"],
            ["Kellele tegevused on suunatud?", "Sagedus", "target_groups"],
            ["Kas seire on täidetud?", "Puuduvate määr", "Andmekvaliteedi mõõdikud"],
            ["Kui pikad on dateeritud tegevused?", "Keskmine kestus", "59 rida alguse ja lõpuga"],
        ],
        col_widths=[6.2, 4.4, 6.0],
    )
    h2(doc, "10.2 Tõlgendust mõjutavad reeglid")
    bullets(
        doc,
        [
            "Keskmine edenemine ei jaga 72 tühja rida nulliks; DAX AVERAGE jätab tühjad välja (~78% kirjeldab ainult 28 rida).",
            "Poliitikameetmete arv aruandes = 105 kirjet, mitte 17 unikaalset meedet.",
            "Partnerluste arv = 176 (summa); unikaalseid partner-id-sid on 39.",
            "Kalendrivaated kasutavad alguskuupäeva; lõppseos on mitteaktiivne.",
        ],
    )
    h2(doc, "10.3 Võrdlusalused")
    add_table(
        doc,
        ["Võrdlus", "Kasutus"],
        [
            ["Sport 2030 (2/3 elanikest liigub)", "Probleemi taust, mitte selle 100 rea KPI"],
            ["WHO GAPPA 4 suunda × 20 meedet", "Katvuse kontroll: millised meetmed on tühjad"],
            ["Staatuse jaotus 19 / 58 / 23", "Portfelli „tervis“"],
            ["Avaliku vaate 100 tegevust", "Ridade arvu kokkusobivus"],
        ],
        col_widths=[7.0, 9.6],
    )

    # ----- 11 -----
    page_break(doc)
    h1(doc, "11. Kirjeldav raport / analüüs")
    h2(doc, "11.1 Portfell tervikuna")
    add_figure(doc, "skeem_staatus.png", "Joonis 7. Tegevuste staatus (n = 100)")
    rich(
        doc,
        [
            "100 tegevusest ",
            ("58 on töös, 23 alustamata, 19 tehtud", True, False),
            " (tehtud osakaal 19,0%). See ei tähenda, et kava on „19% valmis“: 41 tegevust on Pidevad (neist 35 töös) ja kava ulatub 2030. aastani.",
        ],
    )
    add_table(
        doc,
        ["Tüüp", "Pole alustatud", "Töös", "Tehtud", "Kokku"],
        [
            ["Eesti projektid ja programmid", "22", "48", "13", "83"],
            ["Teadus ja innovatsioon", "1", "10", "6", "17"],
            ["Kokku", "23", "58", "19", "100"],
        ],
        col_widths=[5.4, 3.0, 2.6, 2.6, 3.0],
    )
    p(doc, "Teadusest on tehtud 35%, programmidest 16%. Alustamata ridu on peaaegu ainult programmide seas.", align="justify")

    h2(doc, "11.2 WHO suunad")
    add_table(
        doc,
        ["Suund", "Tegevusi", "Alustamata", "Töös", "Tehtud"],
        [
            ["Loome aktiivseid struktuure", "47", "10", "28", "9"],
            ["Loome aktiivset inimest", "22", "1", "16", "5"],
            ["Loome aktiivset keskkonda", "19", "9", "8", "2"],
            ["Loome aktiivset ühiskonda", "12", "3", "6", "3"],
        ],
        col_widths=[6.6, 2.6, 2.6, 2.4, 2.4],
    )
    p(
        doc,
        "Peaaegu pooled tegevused on suunal struktuurid. Aktiivne ühiskond on kõige õhem. Keskkonnasuunal on suhteliselt palju alustamata ridu (9/19).",
        align="justify",
    )

    h2(doc, "11.3 Juhtorganisatsioonid ja võrgustik")
    add_table(
        doc,
        ["Juhtorganisatsioon", "Tegevusi"],
        [
            ["Transpordiamet", "20"],
            ["SA Liikumisharrastuse kompetentsikeskus", "18"],
            ["Kliimaministeerium", "14"],
            ["Kultuuriministeerium", "8"],
            ["Tervise Arengu Instituut", "6"],
            ["HTM, Kaitseressursside Amet, Sotsiaalministeerium", "5 + 5 + 5"],
        ],
        col_widths=[12.4, 4.2],
    )
    p(
        doc,
        "22 unikaalset juhtorganisatsiooni. Transpordi ja kliima osakaal peegeldab rattastrateegia liitumist. Partnerlusi on 176 (78 tegevusel vähemalt üks; keskmiselt ~2,3; maksimum 7). 22 tegevust on ilma partnerita. Unikaalseid partner-id-sid 39.",
        align="justify",
    )

    h2(doc, "11.4 Poliitikameetmed ja sihtrühmad")
    p(
        doc,
        "Tegevused viitavad 17 unikaalsele meetmele 20-st GAPPA meetmest. Meetmekirjeid kokku 105.",
        align="justify",
    )
    add_table(
        doc,
        ["Meede", "Suund", "Tegevusi"],
        [
            ["4.1 Tugevdada poliitikat, juhtimist ja valitsemist", "struktuurid", "16"],
            ["4.2 Parandada ja lõimida andmesüsteeme", "struktuurid", "11"],
            ["4.3 Arendada teadus- ja arendustegevust", "struktuurid", "11"],
            ["1.4 Tugevdada tööjõu suutlikkust", "ühiskond", "10"],
            ["4.5 Arendada uuenduslikke rahastamismehhanisme", "struktuurid", "9"],
            ["3.1 Tugevdada liikumisõpetust ja koolipõhiseid programme", "inimene", "9"],
        ],
        col_widths=[10.2, 3.4, 3.0],
    )
    callout(
        doc,
        "Kolm meedet ilma ühegi tegevuseta",
        "1.2 Kaasnevate hüvede tutvustamine ·  1.3 Pakkuda suuremahulisi liikumisüritusi ·  4.4 Laiendada eestkostetegevust. Need on teadlikkuse, massiürituste ja eestkoste meetmed — kooskõlas õhukese „aktiivse ühiskonna“ suunaga.",
    )
    p(
        doc,
        "Sihtrühmad: Keskkond 31, Lapsed ja noored 30, Teadus ja innovatsioon 21, Elukaareülene 18. Vanemaealistele eraldi silti API-s ei ole (meede 3.4 esineb vaid 1 korral).",
        align="justify",
    )

    h2(doc, "11.5 Aeg, kestus, indikaatorid")
    p(
        doc,
        "Algusaasta (62 dateeritud rida): 2023: 14, 2024: 8, 2025: 17, 2026: 15, 2027: 8, puudub: 38. 59 tegevusel on nii algus kui lõpp; keskmine kestus 1126 päeva (~3,1 aastat), vahemik 364–2191 päeva.",
        align="justify",
    )
    p(
        doc,
        "Indikaatoreid 285 (keskmiselt 2,85 tegevuse kohta). Staatused: Lõpetatud 109 (38,2%), Pooleli 73, Ootel 55, Mõõdik 48. Uuringuviiteid on 161 kirjet 57 tegevusel.",
        align="justify",
    )

    h2(doc, "11.6 Seirekvaliteet")
    add_table(
        doc,
        ["Näitaja", "Arv", "Osakaal"],
        [
            ["Ilma edenemise %-ta", "72", "72%"],
            ["Ilma alguskuupäevata", "38", "38%"],
            ["Ilma partnerita", "22", "22%"],
            ["Tähtaeg ületatud", "0", "0%"],
            ["Edenemine täidetud", "28", "28%"],
            ["Täidetud ridade keskmine edenemine", "78,3%", "ainult need 28"],
        ],
        col_widths=[8.4, 3.6, 4.6],
    )
    p(
        doc,
        "Tähtaegu ei ole ületatud andmestiku reeglite järgi. Edenemise protsent ei ole kasutuskõlblik kogu kava KPI. Õigemad signaalid on staatus, indikaatorite lõpetatus ja puuduvate väärtuste määr.",
        align="justify",
    )

    h2(doc, "11.7 Aruande lehed")
    add_figure(doc, "skeem_aruanne.png", "Joonis 5. Power BI aruande viis lehte")
    add_table(
        doc,
        ["Leht", "Küsimus", "Peamised visuaalid"],
        [
            ["Ülevaade", "Kui palju, mis staatuses, mis suunal?", "KPI-d, tulbad, joon, loend"],
            ["Organisatsioonid", "Kes juhib, kui tihe on võrgustik?", "Juhtide KPI-d, tulbad, rollide tabel"],
            ["Poliitikad", "Millised meetmed ja sihtrühmad?", "Maatriks suund × staatus, tabelid"],
            ["Edenemine", "Kui kaugel ja kui puhas on seire?", "Indikaatorid, tähtaeg, puuduvad väljad"],
            ["Info", "Kuidas mudel töötab?", "INFO.VIEW, leiud, allikad"],
        ],
        col_widths=[3.6, 5.8, 7.2],
    )
    p(
        doc,
        "Ühine navigatsioon ja filtrid: WHO suund, tegevuse tüüp, algusaasta, staatus, juhtorganisatsioon. Nupp Tühista filtrid. Teema LiikumineTheme (taust #F5F5F5, aktsent #2EC4B6).",
        align="justify",
    )

    # ----- 12 -----
    page_break(doc)
    h1(doc, "12. Andmelugu, järeldused")
    callout(
        doc,
        "Lugu ühe lausega",
        "Eesti liikumisaktiivsuse kava on 2026. aasta suveks suur koostööportfell, mis ehitab peamiselt süsteeme — juhtimist, andmeid, teadust ja taristu poliitikat —, samal ajal kui inimeste hoiakute ja massilise osalemise suund on kõige õhem ning seireprotsent on enamiku ridade peal tühi.",
    )
    h2(doc, "12.1 Kolm peatükki")
    rich(
        doc,
        [
            ("1. Kava on elus, mitte „valmis“. ", True, False),
            "58 tegevust töös ja 41 pidevat kirjet on rulluva kava normaalne kuju, mitte mahajäämus. 19 tehtud rida näitavad, et midagi on juba lukku löödud. Sport 2030 siht ei ole selle tabeli ridade arv; see on ühiskonna tulemus, mida need 100 rida peaksid aitama.",
        ],
    )
    rich(
        doc,
        [
            ("2. Raskuskese on struktuuridel, mitte ühiskonnal. ", True, False),
            "47 tegevust suunal „aktiivsed struktuurid“ versus 12 suunal „aktiivne ühiskond“ ei ole juhus. Kolm tühja GAPPA meedet (kaasnevate hüvede tutvustamine, suuremahulised liikumisüritused, eestkoste) on samast perest. Transpordiamet ja Kliimaministeerium juhivad mahult, sest rattastrateegia ja ruum on kavaga liitunud. See on tugevus ja risk: kui struktuurid täituvad, aga kampaaniad jäävad katmata, võib kava jääda asutuste kavaks, mitte inimeste harjumuste kavaks.",
        ],
    )
    rich(
        doc,
        [
            ("3. Seiret saab juhtida staatuse, mitte protsendiga. ", True, False),
            "Edenemise % on täidetud 28 real ja need read on keskmiselt 78% peal — hea, aga valim on kallutatud (kõik 19 tehtut on 100). 72 tühja rida tähendab, et „keskmine edenemine kogu kavas“ oleks vale lugu. Kuupäevad puuduvad pidevatel tegevustel süsteemselt. Tähtaeg ületatud = 0 ütleb, et staatust uuendatakse vähemalt lõpu saabudes. Ausamad signaalid: alustamata ridu keskkonnasuunal (9/19), 22 tegevust ilma partnerita, indikaatorite 38% lõpetatus ja kolm katmata meedet.",
        ],
    )
    h2(doc, "12.2 Soovitused")
    numbered(
        doc,
        [
            "Täita või loobuda edenemise %-st: muuta väli kohustuslikuks Töös ridadel või eemaldada see KPI-st avalikus aruandes.",
            "Pidevatele tegevustele anda vähemalt alguskuupäev (nt kava kinnitamise aasta), et ajatelg ei kaotaks 38% portfellist.",
            "Katta või teadlikult kõrvale jätta meetmed 1.2, 1.3 ja 4.4 — praegu on auk dokumenteerimata.",
            "Normaliseerida partnerid ja meetmed järgmises mudeliversioonis (sildtabelid), kui eesmärk on filtreerida unikaalset partnerit või meedet.",
            "Jälgida keskkonnasuuna alustamata rida — see on ainus suund, kus alustamata on peaaegu sama suur kui töös.",
        ],
    )
    h2(doc, "12.3 Töö piirangud")
    bullets(
        doc,
        [
            "Andmed on hetkeseis, mitte ajajada (last_modified on kitsas aken juunis 2026).",
            "Mitu-mitmele lamedus piirab võrgustikuanalüüsi (nt graaf „kes kellega“).",
            "Sihtrühmade sildid ei kattu 1:1 GAPPA sihtrühmadega (puudu nt eakad eraldi).",
            "Aruanne ei mõõda elanike tegelikku liikumist — ainult kava elluviimise kirjeid.",
            "Grupiliikmete nimed tuleb esilehele lisada enne ametlikku esitamist.",
        ],
    )
    h2(doc, "12.4 Kokkuvõte")
    rich(
        doc,
        [
            "Power BI aruanne ",
            ("Liikumisaktiivsuse programmid", True, False),
            " teeb avalikust API-st loetava seirevaate: 100 tegevust, tähemudel, neli analüüsilehte ja aus andmekvaliteedi vaade. Peamine sõnum juhtidele: ",
            ("kava töötab asutuste ja struktuuride kihil; ühiskonna ja seireprotsendi kiht on veel auklik.", True, False),
            " Järgmine otsus ei ole „kas Power BI töötab“, vaid kas tühjad meetmed ja tühjad edenemised on teadlik valik või seirevõlg.",
        ],
    )

    page_break(doc)
    h1(doc, "Lisad")
    h2(doc, "A. Aruande failid")
    bullets(
        doc,
        [
            "Liikumisaktiivsuse programmid.pbip",
            "Liikumisaktiivsuse programmid.SemanticModel/ (TMDL)",
            "Liikumisaktiivsuse programmid.Report/ (lehed, teema, järjehoidja)",
        ],
    )
    h2(doc, "B. Allikad")
    bullets(
        doc,
        [
            "SA Liikumisharrastuse kompetentsikeskus — liigume.ee/meist, liigume.ee/tegevuskava",
            "Avalik rakendus app.liigume.ee ja API /api/activities, /api/organizations, /api/policies",
            "WHO, Global Action Plan on Physical Activity 2018–2030",
            "Eesti spordipoliitika alusdokument Sport 2030; arengustrateegia Eesti 2035",
            "Avaliku teabe seadus; isikuandmete kaitse üldmäärus (EL) 2016/679",
        ],
    )
    h2(doc, "C. Jooniste loend")
    add_table(
        doc,
        ["Jr", "Joonis"],
        [
            ["1", "Semantiline mudel (tähemudel)"],
            ["2", "Andmeallikad ja andmevoog"],
            ["3", "Grupitöö neli etappi"],
            ["4", "WHO GAPPA suunad tegevuste arvu järgi"],
            ["5", "Power BI aruande viis lehte"],
            ["6", "Andmekvaliteedi neli signaali"],
            ["7", "Tegevuste staatus"],
        ],
        col_widths=[1.6, 15.0],
    )
    p(
        doc,
        "Dokumentatsioon koostatud aruande TMDL-mudeli, aruandelehtede ja avaliku API alusel. Andmete seis 28.08.2026.",
        size=9,
        italic=True,
        color=MUTED,
        space_before=12,
    )

    doc.save(OUT)
    print("Saved", OUT)


if __name__ == "__main__":
    build()
